import plotly.graph_objects as go
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx import Document
import os
import plotly.io as pio
import plotly.express as px
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd

n_rows = 18
df = pd.read_excel(
    f'./transformed_data/modified_rank_prior_{n_rows}.xlsx')
personality_labels = ['Extraversion',	'Agreeableness',
                      'Conscientiousness',	'Neuroticism',	'Openness to Experiences']

mean_threshold = []
std_threshold = []
strong_threshold = []
weak_threshold = []

for idx, (col_name, col_values) in enumerate(df.iloc[:, 4:9].to_dict('list').items()):
    print(col_name)
    strong_threshold.append(df[col_name].mean() + 0.5 * df[col_name].std())
    mean_threshold.append(df[col_name].mean())
    weak_threshold.append(df[col_name].mean() - 0.5 * df[col_name].std())
    std_threshold.append(df[col_name].std())
    pass

print(strong_threshold, weak_threshold)

personalities = {}
for idx, (col_name, col_values) in enumerate(df.iloc[:, 4:9].to_dict('list').items()):
    personalities[f"High {col_name}"] = df[df[col_name]
                                           >= strong_threshold[idx]]
    personalities[f"Low {col_name}"] = df[df[col_name] <= weak_threshold[idx]]

personalities = dict(sorted(personalities.items()))

map_big5_sents = pd.read_csv(
    './data/sentences.csv')
map_big5_sents = map_big5_sents.to_dict('list')
map_big5_sents = dict((vi.lower(), k)
                      for k, v in map_big5_sents.items() for vi in v)
map_big5_sents['Nonverbal communications (e.g. Head nod, gestures, eye contact)'.lower(
)] = 'Voiceless'
map_big5_sents['showing hand gestures and eye contact on the display screen'.lower()
               ] = 'Silent'

Question_answers = [
    ["Good morning, It's nice to see you around", "Hi there!", "Hi there! How's your day going?",
        "Hi, Have a wonderful day!", "Hello", "Nonverbal communications (e.g. Head nod, gestures, eye contact)"],
    ["Please, go ahead. I wouldn't want to hold you up", "If you're in a hurry, please go ahead",
     "Oh, are you going that way? Please keep going", "Hey there! No need to stand on ceremony, go right ahead",
     "Please, go ahead", "Nonverbal communications (e.g. Head nod, gestures, eye contact)"],
    ["That's very kind of you, I appreciate it. Have a nice day!", "That's very kind of you, thank you", "Hey, thanks a bunch for letting me go ahead!",
        "Thank you! Your help brightened my day", "Thank you", "Nonverbal communications (e.g. head nod, gestures, eye contact)"],
    ["Good morning, It's nice to see you around", "Hi there!", "Hi there! How's your day going?",
        "Hi, Have a wonderful day!", "Hello", 'Showing hand gestures and eye contact on the display screen'],
    ["Please, go ahead. I wouldn't want to hold you up", "If you're in a hurry, please go ahead", "Oh, are you going that way? Please keep going",
        "Hey there! No need to stand on ceremony, go right ahead", "Please, go ahead", 'Showing hand gestures and eye contact on the display screen'],
    ["That's very kind of you, I appreciate it. Have a nice day!", "That's very kind of you, thank you", "Hey, thanks a bunch for letting me go ahead!",
        "Thank you! Your help brightened my day", "Thank you", 'Showing hand gestures and eye contact on the display screen']
]

# Sample data


def plot_spider_graph(mean, std, threshold_mean, threshold_std, save_name, title):
    big5_labels = ['Extraversion',	'Agreeableness',
                   'Conscientiousness',	'Neuroticism',	'Openness to Experience']

    # mean_sub_std = [round(mean[i] - std[i], 3) for i in range(len(mean))]

    # mean_add_std = [round(mean[i] + std[i], 3)  for i in range(len(mean))]

    threshold_mean = [round(el, 3) for el in threshold_mean]

    threshold_std = [round(el, 3) for el in threshold_std]

    values = mean + [mean[i] - std[i]
                     for i in range(len(mean))] + [mean[i] + std[i] for i in range(len(mean))]

    values = [round(vi, 3) for vi in values]

    df = pd.DataFrame(dict(
        value=values,
        variable=big5_labels*3,
        group=['mean', 'mean', 'mean', 'mean', 'mean',
               '-std', '-std', '-std', '-std', '-std',
               '+std', '+std', '+std', '+std', '+std'],
        name='Result'))

    # df_low = pd.DataFrame(dict(
    #     value=threshold_mean - 0.5*threshold_std,
    #     variable=big5_labels,
    #     group=['low', 'low', 'low', 'low', 'low'],
    #     name='Lower Threshold'
    # ))

    # df_high = pd.DataFrame(dict(
    #     value=threshold_mean + 0.5*threshold_std,
    #     variable=big5_labels,
    #     group=['high', 'high', 'high', 'high', 'high'],
    #     name='Upper Threshold',

    # ))

    df_mean = pd.DataFrame(dict(
        value=threshold_mean,
        variable=big5_labels,
        group=['mean', 'mean', 'mean', 'mean', 'mean'],
        name='Mean population',

    ))

    df = pd.concat([df, df_mean], axis=0, ignore_index=True)

    fig = px.line_polar(df, r='value', theta='variable',
                        line_close=True, template='seaborn',
                        color='name',
                        line_dash='group')

    fig.update_traces(textposition='top center')

    fig.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,
                range=[0, 15]
            )),
        showlegend=True,
        font={'size': 12},
        title={'text': f'<b>{title}</b>', 'font': {'size': 12}}
    )

    scale = 0.8
    os.makedirs('./figures/survey_spider_images/', exist_ok=True)
    pio.write_image(
        fig, f"./figures/survey_spider_images/{save_name}.png", width=480/scale, height=300/scale, scale=scale)


# ####
document = Document()
p = document.add_paragraph()
r = p.add_run()

df_list = []
for idx, (col_name, col_values) in enumerate(df.iloc[:, 9:].to_dict('List').items()):
    if 'Answer_0' in col_name:
        for uniq_val in Question_answers[idx % 6]:
            if type(uniq_val) != str:
                continue
            # if uniq_val.lower() not in map_big5_sents.keys():
            #     continue

            df_selected = df.loc[df[col_name] == uniq_val].iloc[:, 4:9]
            # if len(df_selected) == 0:
            #     print(df[col_name], '\n', uniq_val)
            big5_mean = list(df_selected.mean().to_dict().values())
            big5_std = list(df_selected.std().to_dict().values())

            plot_spider_graph(big5_mean, big5_std, mean_threshold, std_threshold,
                              f"{col_name}_content_{uniq_val.replace('/', '')}",
                              title=f"{uniq_val} ({map_big5_sents[uniq_val.lower()]}) ({len(df_selected)})")

            r.add_text(f">>{col_name}: ")
            r.add_picture(
                f"./figures/survey_spider_images/{col_name}_content_{uniq_val.replace('/', '')}.png")

document.save(f'./plot_documents/survey_spider_charts_{n_rows}.docx')
