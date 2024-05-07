import plotly.graph_objects as go
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx import Document
import os
import plotly.io as pio
import plotly.express as px
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd


map_big5_sents = pd.read_csv(
    './data/sentences.csv')
map_big5_sents = map_big5_sents.to_dict('list')
map_big5_sents = dict((vi.lower(), k)
                      for k, v in map_big5_sents.items() for vi in v)
map_big5_sents['Nonverbal communications (e.g. Head nod, gestures, eye contact)'.lower(
)] = '-'
map_big5_sents['showing hand gestures and eye contact on the display screen'.lower()
               ] = '-'

# map_big5_sents['-'] = '-'
key_order = ['Agreeableness', 'Conscientiousness',
             'Openness',	'Extraversion',	'Neuroticism', '-']
ques_ans_map = {}
for idx, (k, v) in enumerate(map_big5_sents.items()):
    ques_ans_map.setdefault(idx % 3, set()).add(k)

ques_ans_map[0].add(
    'Nonverbal communications (e.g. Head nod, gestures, eye contact)'.lower())
ques_ans_map[1].add(
    'Nonverbal communications (e.g. Head nod, gestures, eye contact)'.lower())
ques_ans_map[2].add(
    'Nonverbal communications (e.g. Head nod, gestures, eye contact)'.lower())
ques_ans_map[0].add(
    'showing hand gestures and eye contact on the display screen'.lower())
ques_ans_map[1].add(
    'showing hand gestures and eye contact on the display screen'.lower())
ques_ans_map[2].add(
    'showing hand gestures and eye contact on the display screen'.lower())


def sublist(lst1, lst2):
    return set(lst1) <= set(lst2)


n_rows = 18
df = pd.read_excel(
    f'./transformed_data/modified_sent_prior_{n_rows}.xlsx')

strong_threshold = []
weak_threshold = []

for idx, (col_name, col_values) in enumerate(df.iloc[:, 4:9].to_dict('list').items()):
    # print(col_name)
    strong_threshold.append(df[col_name].mean() + 0.5 * df[col_name].std())
    weak_threshold.append(df[col_name].mean() - 0.5 * df[col_name].std())
    pass

# print(strong_threshold, weak_threshold)

personalities = {}
for idx, (col_name, col_values) in enumerate(df.iloc[:, 4:9].to_dict('list').items()):
    personalities[f"High {col_name}"] = df[df[col_name]
                                           >= strong_threshold[idx]]
    personalities[f"Low {col_name}"] = df[df[col_name] <= weak_threshold[idx]]

personalities = dict(sorted(personalities.items()))

document = Document()
p = document.add_paragraph()
r = p.add_run()

for personality, personality_df in personalities.items():
    # persnality label and the dataframe filtered by the label
    # print(personality_df)
    for i in range(6):
        df_selected = personality_df.iloc[:, 9 + i * 6: 9 + i * 6 + 6]
        # print(df_selected)
        situation, question, answer_ = list(df_selected.columns)[0].split('-')
        # print(df_selected.columns)
        myKeys = list([x.split('-')[-1] for x in df_selected.keys()])
        myKeys.sort()
        sorted_dict = {'-'.join([situation, question, i]): df_selected['-'.join([situation, question, i])] for i in myKeys}

        data = [col_list for col_name, col_list in sorted_dict.items()]
        labels = [col_name.split('-')[-1][5:-1]
                  for col_name, col_list in sorted_dict.items()]
        # print(data)
        fig = plt.figure(figsize=(3, 1.8))
        # Creating axes instance
        ax = fig.add_axes([0, 0, 1, 1])
        ax.set_yticklabels(labels, fontsize=8)
        # Creating plot
        bp = ax.boxplot(data, vert=False, patch_artist=True)

        colors = ['pink', 'lightblue', 'lightgreen', 'lightyellow', 'purple']

        for patch, color in zip(bp['boxes'], colors):
            patch.set_facecolor(color)

        plt.tight_layout()
        # plt.grid(True)
        # plt.show()
        os.makedirs('./figures/survey_box_images/', exist_ok=True)
        plt.savefig(
            f'./figures/survey_box_images/{personality}-{situation}-{question}.png', bbox_inches='tight')
        plt.close()

        r.add_text(f">>{personality} in {situation} with {question}:")
        r.add_picture(
            f"./figures/survey_box_images/{personality}-{situation}-{question}.png")

document.save(
    './plot_document/box_plot_personality.docx')

# ///////////////////////////////////////////////////////////////
# sentences_df = {}
# for idx, (col_name, col_values) in enumerate(df.iloc[:, 9:].to_dict('list').items()):
#     sentences_df[f"{col_name}"] = df[df[col_name] == 6].iloc[:, 4: 9]

# sentences_df = dict(sorted(sentences_df.items()))

# document = Document()
# p = document.add_paragraph()
# r = p.add_run()

# for i, (sent, sent_df) in enumerate(sentences_df.items()):
#     # persnality label and the dataframe filtered by the label
#     # print(personality_df)
#     situation, question, sent_ = sent.split('-')
#     myKeys = list(sent_df.keys())
#     myKeys.sort()
#     sorted_dict = {i: sent_df[i] for i in myKeys}

#     data = [col_list for col_name, col_list in sorted_dict.items()]
#     labels = [col_name for col_name, col_list in sorted_dict.items()]
#     # print(data)
#     fig = plt.figure(figsize=(4.5, 2))
#     # Creating axes instance
#     ax = fig.add_axes([0, 0, 1, 1])
#     ax.set_yticklabels(labels, fontsize=8)
#     # Creating plot
#     bp = ax.boxplot(data, vert=False, patch_artist=True)

#     ax.set_title(sent_,  fontsize=12)
#     ax.set_xlim(0, 14)

#     colors = ['pink', 'lightblue', 'lightgreen', 'lightyellow', 'purple']

#     for patch, color in zip(bp['boxes'], colors):
#         patch.set_facecolor(color)

#     plt.tight_layout()
#     # plt.grid(True)
#     # plt.show()

#     os.makedirs('./figures/survey_box_images2/', exist_ok=True)
#     plt.savefig(
#         f'./figures/survey_box_images2/sent{i}-{situation}-{question}.png', bbox_inches='tight')
#     plt.close()

#     r.add_text(f">>{situation} with {question} for {sent_}:")
#     r.add_picture(
#         f"./figures/survey_box_images2/sent{i}-{situation}-{question}.png")

# document.save(
#     './plot_document/box_plot_sentence_rank_20.docx')
