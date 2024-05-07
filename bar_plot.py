import plotly.graph_objects as go
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx import Document
import os
import plotly.io as pio
import plotly.express as px
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd

n_rows = 18
df = pd.read_excel(
    f'./transformed_data/modified_rank_prior_{n_rows}.xlsx')
personality_labels = ['Extraversion',	'Agreeableness',
                      'Conscientiousness',	'Neuroticism',	'Openness to Experiences']

strong_threshold = []
weak_threshold = []

for idx, (col_name, col_values) in enumerate(df.iloc[:, 4:9].to_dict('list').items()):
    print(col_name)
    strong_threshold.append(df[col_name].mean() + 0.5 * df[col_name].std())
    weak_threshold.append(df[col_name].mean() - 0.5 * df[col_name].std())
    pass

print(strong_threshold, weak_threshold)

personalities = {}
for idx, (col_name, col_values) in enumerate(df.iloc[:, 4:9].to_dict('list').items()):
    personalities[f"High {col_name}"] = df[df[col_name]
                                           >= strong_threshold[idx]]
    personalities[f"Low {col_name}"] = df[df[col_name] <= weak_threshold[idx]]

personalities = dict(sorted(personalities.items()))

map_big5_sents = pd.read_csv(
    './data/sentences.csv')
map_big5_sents = map_big5_sents.to_dict('list')
map_big5_sents = dict((vi.lower(), k)
                      for k, v in map_big5_sents.items() for vi in v)
map_big5_sents['Nonverbal communications (e.g. Head nod, gestures, eye contact)'.lower(
)] = '-'
map_big5_sents['showing hand gestures and eye contact on the display screen'.lower()
               ] = '-'

# map_big5_sents['-'] = '-'
key_order = ['Agreeableness', 'Conscientiousness',
             'Openness',	'Extraversion',	'Neuroticism', '-']
ques_ans_map = {}
for idx, (k, v) in enumerate(map_big5_sents.items()):
    ques_ans_map.setdefault(idx % 3, set()).add(k)
ques_ans_map[0].add(
    'Nonverbal communications (e.g. Head nod, gestures, eye contact)'.lower())
ques_ans_map[1].add(
    'Nonverbal communications (e.g. Head nod, gestures, eye contact)'.lower())
ques_ans_map[2].add(
    'Nonverbal communications (e.g. Head nod, gestures, eye contact)'.lower())
ques_ans_map[0].add(
    'showing hand gestures and eye contact on the display screen'.lower())
ques_ans_map[1].add(
    'showing hand gestures and eye contact on the display screen'.lower())
ques_ans_map[2].add(
    'showing hand gestures and eye contact on the display screen'.lower())


def sublist(lst1, lst2):
    return set(lst1) <= set(lst2)

# consider 1st answers


document = Document()
p = document.add_paragraph()
r = p.add_run()

df_list = []
for idx, (col_name, col_values) in enumerate(df.iloc[:, 9:].to_dict('list').items()):
    if 'Answer_0' in col_name:
        situation, question, answer = col_name.split('-')
        for personality, p_dict in (personalities.items()):
            D = p_dict[col_name].value_counts().to_dict()
            # answer count appearance
            D.pop('-', None)

            # NOTE: sort dict
            # fill non value with 0
            D_lower = {k.lower(): v for k, v in D.items()}
            for full_list in ques_ans_map.values():
                if sublist(D_lower.keys(), full_list):
                    for k in set(full_list) - set(D_lower.keys()):
                        D[k.capitalize()] = 0

            D_sorted = {}
            for key in key_order:  # sort sentences in v 5 personalities
                for sent, count in D.items():
                    if sent.lower() != "please, go ahead. i wouldn't want to hold you up.":
                        if map_big5_sents[sent.lower()] == key:
                            D_sorted[f"{sent} [{map_big5_sents[sent.lower()]}]"] = count
                            continue
                    else:
                        if map_big5_sents[sent.lower()[:-1]] == key:
                            D_sorted[f"{sent} [{map_big5_sents[sent.lower()[:-1]]}]"] = count
                            continue
            # sort keys
            # D_sorted = dict(sorted(D_sorted.items()))
            # print(personality, D_sorted)

            # plot
            plt.figure(figsize=(6, 1.5), dpi=100)
            plt.barh(range(len(D_sorted)), list(
                D_sorted.values()), align='center')

            plt.yticks(range(len(D_sorted)), list(D_sorted.keys()), fontsize=8)

            max_x = 10
            plt.xlim(0, max_x)
            plt.xticks(np.arange(0, max_x-1, step=1))

            plt.gca().invert_yaxis()
            plt.tight_layout()

            os.makedirs('./figures/survey_bar_images/', exist_ok=True)
            plt.savefig(
                f'./figures/survey_bar_images/{personality}-{situation}-{question}.png')
            plt.close()

            r.add_text(f">>{personality} in {situation} with {question}:")
            r.add_picture(
                f"./figures/survey_bar_images/{personality}-{situation}-{question}.png")
            # r.add_text(f"\n\n\n")

document.save(
    f'./plot_document/survey_bar_charts_{n_rows}.docx')

# stacked bar
document = Document()
p = document.add_paragraph()
r = p.add_run()

document1 = Document()
p1 = document1.add_paragraph()
r1 = p1.add_run()

df_list = {}
personality_stackedbar_data = {}
for idx, (col_name, col_values) in enumerate(df.iloc[:, 9:].to_dict('list').items()):
    situation, question, answer = col_name.split('-')

    rank = answer.split('_')[-1]
    if rank == 'prefered':
        continue

    df_list.setdefault(question, {}).setdefault(rank, [])
    for personality, p_dict in (personalities.items()):
        D = p_dict[col_name].value_counts().to_dict()
        # answer count appearance
        D.pop('-', None)
        for sent, sent_count in D.items():
            personality_stackedbar_data.setdefault(situation, {}).setdefault(question, {}).setdefault(
                personality, {}).setdefault(rank, {}).setdefault(sent, sent_count)

for situation, question_data in personality_stackedbar_data.items():
    for question, answer_data in question_data.items():
        for personality, rank_data in answer_data.items():
            plot_data = {}
            plot_keys = []
            for rank, data in rank_data.items():
                # fill non value with 0
                D_lower = {sent.lower(): sent_count for sent,
                           sent_count in data.items()}
                for full_list in ques_ans_map.values():
                    if sublist(D_lower.keys(), full_list):
                        for sent in set(full_list) - set(D_lower.keys()):
                            data[sent.capitalize()] = 0
                # ....
                if situation == 'Case_0':
                    data.pop(
                        'Showing hand gestures and eye contact on the display screen', None)
                elif situation == 'Case_1':
                    data.pop(
                        'Nonverbal communications (e.g. head nod, gestures, eye contact)', None)

                data_sorted = {}
                # sort sentences in 5 personalities with form "sent[personaltiy]"
                for key in key_order:
                    for sent, count in data.items():
                        if sent.lower() != "please, go ahead. i wouldn't want to hold you up.":
                            if map_big5_sents[sent.lower()] == key:
                                data_sorted[f"{sent} [{map_big5_sents[sent.lower()]}]"] = count
                                continue
                        else:
                            if map_big5_sents[sent.lower()[:-1]] == key:
                                data_sorted[f"{sent} [{map_big5_sents[sent.lower()[:-1]]}]"] = count
                                continue
                # print(rank, data_sorted)
                plot_data.setdefault(rank, list(data_sorted.values()))
                plot_keys = list(data_sorted.keys())

            plot_data_df = pd.DataFrame.from_dict(plot_data)
            plot_data_df['Sentence'] = plot_keys
            plot_data_df_norm = plot_data_df.iloc[:,
                                                  :-1].apply(lambda x: x / sum(x), axis=1)
            plot_data_df_norm['Sentence'] = plot_keys
            plot = plot_data_df_norm.plot.barh(
                x='Sentence', stacked=True, mark_right=True, legend=False)

            # # plot
            fig = plot.get_figure()
            plt.rcParams['figure.figsize'] = [6, 1.5]
            plt.yticks(range(len(plot_keys)), list(plot_keys), fontsize=6)

            plt.gca().invert_yaxis()
            plt.tight_layout()

            os.makedirs('./figures/survey_bar_images/', exist_ok=True)
            fig.savefig(
                f'./figures/survey_bar_images/{personality}-{situation}-{question}_norm.png')
            plt.close()

            r.add_text(f">>{personality} in {situation} with {question}:")
            r.add_picture(
                f"./figures/survey_bar_images/{personality}-{situation}-{question}_norm.png")

            plot1 = plot_data_df.plot.barh(
                x='Sentence', stacked=True, mark_right=True, legend=False)
            fig1 = plot1.get_figure()
            plt.rcParams['figure.figsize'] = [6, 1.5]
            plt.yticks(range(len(plot_keys)), list(plot_keys), fontsize=6)
            plt.gca().invert_yaxis()
            plt.tight_layout()
            fig1.savefig(
                f'./figures/survey_bar_images/{personality}-{situation}-{question}.png')
            plt.close()
            r1.add_text(f">>{personality} in {situation} with {question}:")
            r1.add_picture(
                f"./figures/survey_bar_images/{personality}-{situation}-{question}.png")

            # r.add_text(f"\n\n\n")
    # break

document.save(
    f'./plot_document/survey_bar_charts_norm_{n_rows}.docx')
document1.save(
    f'./plot_document/survey_bar_charts_{n_rows}.docx')
