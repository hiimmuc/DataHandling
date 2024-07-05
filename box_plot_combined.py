import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib.patches import Patch

import plotly.graph_objects as go
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx import Document
import os
import plotly.io as pio
import plotly.express as px
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
from tqdm import tqdm


def sublist(lst1, lst2):
    return set(lst1) <= set(lst2)


# document = Document()
# p = document.add_paragraph()
# r = p.add_run()

# n_rows = 18
# df = pd.read_excel(
#     f'./transformed_data/modified_sent_prior_{n_rows}_avg.xlsx')

# key_order = ['Agreeableness', 'Conscientiousness',
#              'Openness to Experiences',	'Extraversion',	'Neuroticism', 'Voiceless', 'Silent']

# std_scaling_factor = 0.5

# strong_threshold = []
# weak_threshold = []

# for idx, (col_name, col_values) in enumerate(df.iloc[:, 4:9].to_dict('list').items()):
#     strong_threshold.append(df[col_name].mean() +
#                             std_scaling_factor * df[col_name].std())
#     weak_threshold.append(df[col_name].mean() -
#                           std_scaling_factor * df[col_name].std())


# print(strong_threshold, weak_threshold)

# personality_dataFrames = {'High': {}, 'Low': {}}


# for idx, (col_name, col_values) in enumerate(df.iloc[:, 4:9].to_dict('list').items()):

#     personality_dataFrames['High'][col_name] = df[df[col_name]
#                                                   >= strong_threshold[idx]]
#     personality_dataFrames['Low'][col_name] = df[df[col_name]
#                                                  <= weak_threshold[idx]]

# # for idx, (col_name, col_values) in enumerate(df.iloc[:, 4:9].to_dict('list').items()):

# #     personality_dataFrames['High'][col_name] = df[df[col_name]
# #                                                   >= 7.0]
# #     personality_dataFrames['Low'][col_name] = df[df[col_name]
# #                                                  < 7.0]

# # personality_dataFrames = dict(sorted(personality_dataFrames.items()))
# personality_labels = ['Extraversion', 'Agreeableness', 'Conscientiousness',
#                       'Neuroticism', 'Openness to Experiences']

# # print(personalities)

# # Settings
# x = 6  # Want figures to be A6
# plt.rc('figure', figsize=[46.82 * .5**(.5 * x), 33.11 * .5**(.5 * x)])
# plt.rc('text', usetex=True)
# plt.rc('font', family='serif')

# # Define which colours you want to use
# colors = ['blue', 'red', 'lightblue', 'pink']
# # Define the groups
# groups = ['High score of personality',
#           'Low score of personality',]

# dataFrames_dict = [personality_dataFrames[level]
#                    for level in ['High', 'Low']]

# # Get the max of the dataset
# y_max = 18
# # Get the min of the dataset
# y_min = 0
# # Calculate the y-axis range
# y_range = y_max - y_min


# # print("", len(x_pos_range))
# # # Plot
# df_template = None
# for case in range(2):  # case 0, 1
#     for personality in tqdm(personality_labels):  # personality 5
#         # generate data for box plot
#         datasets = []
#         for idx, df_list in enumerate(dataFrames_dict):  # 4 each label in col
#             # select cols
#             start_index = 9 + case * 6
#             df_qa = df_list[personality].iloc[:, start_index: start_index + 6]
#             df_qa = df_qa.reset_index(drop=True)

#             datasets.append(df_qa)

#         assert len(datasets) == 2
#         # # Create the plot
#         fig = plt.figure(f"Case_{case}-{personality}",  figsize=(11.7, 8.3))

#         ax = plt.axes()
#         # Set x-positions for boxes
#         x_pos_range = np.arange(len(datasets)) / (len(datasets) - 1)
#         x_pos = (x_pos_range * 0.5) + 0.75

#         for i, data in enumerate(datasets):
#             positions = [x_pos[i] + j * 1 for j in range(len(data.T))]
#             bp = ax.boxplot(
#                 np.array(data), sym='', whis=[0, 100], widths=0.6 / len(datasets),
#                 labels=list(x.split('with')[-1] for x in datasets[0]), patch_artist=True,
#                 positions=positions
#             )
#             # Fill the boxes with colours (requires patch_artist=True)
#             k = i % len(colors)
#             for box in bp['boxes']:
#                 box.set(facecolor=colors[k])
#             # Make the median lines more visible
#             plt.setp(bp['medians'], color='black')

#             # Get the samples' medians
#             medians = [bp['medians'][j].get_ydata()[0]
#                        for j in range(len(data.T))]
#             medians = [str(round(s, 1)) for s in medians]
#             # Increase the height of the plot by 5% to fit the labels
#             ax.set_ylim([y_min - 0.1 * y_range, y_max + 0.05 * y_range])
#             # Set the y-positions for the labels
#             y_pos = y_min - 0.075 * y_range
#             for tick, label in zip(range(len(data.T)), ax.get_xticklabels()):
#                 k = tick % 2
#                 ax.text(
#                     positions[tick], y_pos,
#                     f'{medians[tick]}',
#                     horizontalalignment='center', size='large'
#                 )
#         # Axis details
#         details = ax.set(
#             title=f"Case {case} with participants with high / low scores of {personality}",
#             ylabel='score'
#         )

#         ax.set_xlabel('Sentence sets')
#         ax.tick_params(axis='x', bottom=False)

#         xticks = ax.set_xticks(np.arange(len(list(datasets[0]))) + 1)
#         xticks = ax.set_xticks(
#             np.arange(len(list(datasets[0])) + 1) + 0.5, minor=True)
#         xlim = ax.set_xlim([0.5, len(list(datasets[0])) + 0.5])
#         # Legend
#         legend_elements = []
#         for i in range(len(datasets)):
#             j = i % len(groups)
#             k = i % len(colors)
#             legend_elements.append(Patch(facecolor=colors[k], label=groups[j]))
#         # ax.legend(handles=legend_elements, fontsize=8)
#         plt.gca().legend(
#             legend_elements, groups,
#             fontsize=8, loc='center left', bbox_to_anchor=(1, 0.5)
#         )
#         # Background
#         # ax.set_facecolor('0.8')
#         plt.grid(True, color='black', linestyle='-.', linewidth=0.2)
#         # plt.show()
#         # plt.tight_layout()

#         os.makedirs('./figures/survey_box_images_avg/', exist_ok=True)
#         plt.savefig(
#             f'./figures/survey_box_images_avg/Case_{case}-{personality}.png', bbox_inches='tight')
#         plt.close()

#         # add to document
#         r.add_text(f">>{personality} in Case_{case}:")
#         r.add_picture(
#             f"./figures/survey_box_images_avg/Case_{case}-{personality}.png")

# document.save(
#     './plot_document/box_plot_personality_avg.docx')

# /////////////////////////////////////////////////////////////////

# document = Document()
# p = document.add_paragraph()
# r = p.add_run()

# n_rows = 18
# df = pd.read_excel(
#     f'./transformed_data/modified_sent_prior_{n_rows}_avg.xlsx')

# key_order = ['Agreeableness', 'Conscientiousness',
#              'Openness to Experiences',	'Extraversion',	'Neuroticism', 'NonVerbal', 'Silent']


# df_vn = df[df['What is your nationality?']
#            == 'Vietnamese']
# df_jp = df[df['What is your nationality?']
#            == 'Japanese']

# strong_threshold_vn = []
# strong_threshold_jp = []

# weak_threshold_vn = []
# weak_threshold_jp = []


# for idx, (col_name, col_values) in enumerate(df_vn.iloc[:, 4:9].to_dict('list').items()):
#     strong_threshold_vn.append(df_vn[col_name].mean(
#     ) + std_scaling_factor * df_vn[col_name].std())
#     weak_threshold_vn.append(df_vn[col_name].mean() -
#                              std_scaling_factor * df_vn[col_name].std())
# print(strong_threshold_vn, weak_threshold_vn)

# for idx, (col_name, col_values) in enumerate(df_jp.iloc[:, 4:9].to_dict('list').items()):
#     strong_threshold_jp.append(df_jp[col_name].mean(
#     ) + std_scaling_factor * df_jp[col_name].std())
#     weak_threshold_jp.append(df_jp[col_name].mean() -
#                              std_scaling_factor * df_jp[col_name].std())
# print(strong_threshold_jp, weak_threshold_jp)

# personality_dataFrames = {'Vietnamese': {'High': {}, 'Low': {}},
#                           'Japanese': {'High': {}, 'Low': {}}}

# # for idx, (col_name, col_values) in enumerate(df.iloc[:, 4:9].to_dict('list').items()):

# #     personality_dataFrames['Vietnamese']['High'][col_name] = df_vn[df_vn[col_name]
# #                                                                    >= strong_threshold_vn[idx]]
# #     personality_dataFrames['Vietnamese']['Low'][col_name] = df_vn[df_vn[col_name]
# #                                                                   <= weak_threshold_vn[idx]]
# #     personality_dataFrames['Japanese']['High'][col_name] = df_jp[df_jp[col_name]
# #                                                                  >= strong_threshold_jp[idx]]
# #     personality_dataFrames['Japanese']['Low'][col_name] = df_jp[df_jp[col_name]
# #                                                                 <= weak_threshold_jp[idx]]


# for idx, (col_name, col_values) in enumerate(df.iloc[:, 4:9].to_dict('list').items()):
#     strong_avg = (strong_threshold_vn[idx] + strong_threshold_jp[idx]) / 2
#     weak_avg = (weak_threshold_vn[idx] + weak_threshold_jp[idx]) / 2

#     personality_dataFrames['Vietnamese']['High'][col_name] = df_vn[df_vn[col_name]
#                                                                    >= strong_avg]
#     personality_dataFrames['Vietnamese']['Low'][col_name] = df_vn[df_vn[col_name]
#                                                                   <= weak_avg]
#     personality_dataFrames['Japanese']['High'][col_name] = df_jp[df_jp[col_name]
#                                                                  >= strong_avg]
#     personality_dataFrames['Japanese']['Low'][col_name] = df_jp[df_jp[col_name]
#                                                                 <= weak_avg]

# # personality_dataFrames = dict(sorted(personality_dataFrames.items()))
# personality_labels = ['Extraversion', 'Agreeableness', 'Conscientiousness',
#                       'Neuroticism', 'Openness to Experiences']

# # print(personalities)

# # Settings
# x = 6  # Want figures to be A6
# plt.rc('figure', figsize=[46.82 * .5**(.5 * x), 33.11 * .5**(.5 * x)])
# plt.rc('text', usetex=True)
# plt.rc('font', family='serif')

# # Define which colours you want to use
# colors = ['blue', 'red', 'lightblue', 'pink']
# # Define the groups
# groups = ['High score of Vietnamese',
#           'Low score of Vietnamese',
#           'High score of Japanese',
#           'Low score of Japanese']

# dataFrames_dict = [personality_dataFrames[nationality][level]
#                    for nationality in personality_dataFrames.keys()
#                    for level in ['High', 'Low']]

# # Get the max of the dataset
# y_max = 18
# # Get the min of the dataset
# y_min = 0
# # Calculate the y-axis range
# y_range = y_max - y_min


# # print("", len(x_pos_range))
# # # Plot
# df_template = None
# for case in range(2):  # case 0, 1
#     for personality in tqdm(personality_labels):  # personality 5
#         # generate data for box plot
#         datasets = []
#         for idx, df_list in enumerate(dataFrames_dict):  # 4 each label in col
#             # select cols
#             start_index = 9 + case * 6
#             df_qa = df_list[personality].iloc[:, start_index: start_index + 6]
#             df_qa = df_qa.reset_index(drop=True)

#             datasets.append(df_qa)

#         assert len(datasets) == 4
#         # # Create the plot
#         fig = plt.figure(f"Case_{case}-{personality}", figsize=(11.7, 8.3))
#         ax = plt.axes()
#         # Set x-positions for boxes
#         x_pos_range = np.arange(len(datasets)) / (len(datasets) - 1)
#         x_pos = (x_pos_range * 0.5) + 0.75

#         for i, data in enumerate(datasets):
#             positions = [x_pos[i] + j * 1 for j in range(len(data.T))]
#             bp = ax.boxplot(
#                 np.array(data), sym='', whis=[0, 100], widths=0.6 / len(datasets),
#                 labels=list(x.split('with')[-1] for x in datasets[0]), patch_artist=True,
#                 positions=positions
#             )
#             # Fill the boxes with colours (requires patch_artist=True)
#             k = i % len(colors)
#             for box in bp['boxes']:
#                 box.set(facecolor=colors[k])
#             # Make the median lines more visible
#             plt.setp(bp['medians'], color='black')

#             # Get the samples' medians
#             medians = [bp['medians'][j].get_ydata()[0]
#                        for j in range(len(data.T))]
#             medians = [str(round(s, 1)) for s in medians]
#             # Increase the height of the plot by 5% to fit the labels
#             ax.set_ylim([y_min - 0.1 * y_range, y_max + 0.05 * y_range])
#             # Set the y-positions for the labels
#             y_pos = y_min - 0.075 * y_range
#             for tick, label in zip(range(len(data.T)), ax.get_xticklabels()):
#                 k = tick % 2
#                 ax.text(
#                     positions[tick], y_pos,
#                     f'{medians[tick]}',
#                     horizontalalignment='center', size='medium'
#                 )
#         # Axis details
#         details = ax.set(
#             title=f"Case {case} with participants with high / low scores of {personality}",
#             ylabel='score'
#         )

#         ax.set_xlabel('Sentence sets')
#         ax.tick_params(axis='x', bottom=False)

#         xticks = ax.set_xticks(np.arange(len(list(datasets[0]))) + 1)
#         xticks = ax.set_xticks(
#             np.arange(len(list(datasets[0])) + 1) + 0.5, minor=True)
#         xlim = ax.set_xlim([0.5, len(list(datasets[0])) + 0.5])
#         # Legend
#         legend_elements = []
#         for i in range(len(datasets)):
#             j = i % len(groups)
#             k = i % len(colors)
#             legend_elements.append(Patch(facecolor=colors[k], label=groups[j]))
#         # ax.legend(handles=legend_elements, fontsize=8)
#         plt.gca().legend(
#             legend_elements, groups,
#             fontsize=8, loc='center left', bbox_to_anchor=(1, 0.5)
#         )
#         # Background
#         # ax.set_facecolor('0.8')
#         plt.grid(True, color='black', linestyle='-.', linewidth=0.2)

#         plt.tight_layout()
#         # plt.show()

#         # save
#         os.makedirs('./figures/survey_box_images_avg_culture/', exist_ok=True)
#         plt.savefig(
#             f'./figures/survey_box_images_avg_culture/Case_{case}-{personality}.png', bbox_inches='tight')
#         plt.close()

#         # add to document
#         r.add_text(f">>{personality} in Case_{case}:")
#         r.add_picture(
#             f"./figures/survey_box_images_avg_culture/Case_{case}-{personality}.png")

# document.save(
#     './plot_document/box_plot_personality_avg_culture.docx')


# +++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
document = Document()

p = document.add_paragraph()
r = p.add_run()


df = pd.read_excel(
    f'./transformed_data/onvehicleexp.xlsx')

key_order = ['Agreeableness', 'Conscientiousness',
             'Openness to Experiences',	'Extraversion',	'Neuroticism', 'Silent']

std_scaling_factor = 0

strong_threshold = []
weak_threshold = []

for idx, (col_name, col_values) in enumerate(df.iloc[:, 6:6+5].to_dict('list').items()):
    strong_threshold.append(df[col_name].mean() +
                            std_scaling_factor * df[col_name].std())
    weak_threshold.append(df[col_name].mean() -
                          std_scaling_factor * df[col_name].std())


print(strong_threshold, weak_threshold)

personality_dataFrames = {'High': {}, 'Low': {}}


for idx, (col_name, col_values) in enumerate(df.iloc[:, 6:6+5].to_dict('list').items()):

    personality_dataFrames['High'][col_name] = df[df[col_name]
                                                  >= strong_threshold[idx]]
    personality_dataFrames['Low'][col_name] = df[df[col_name]
                                                 <= weak_threshold[idx]]

# for idx, (col_name, col_values) in enumerate(df.iloc[:, 4:9].to_dict('list').items()):

#     personality_dataFrames['High'][col_name] = df[df[col_name]
#                                                   >= 7.0]
#     personality_dataFrames['Low'][col_name] = df[df[col_name]
#                                                  < 7.0]

# personality_dataFrames = dict(sorted(personality_dataFrames.items()))
personality_labels = ['Extraversion', 'Agreeableness', 'Conscientiousness',
                      'Neuroticism', 'Openness to Experiences']

# print(personalities)

# Settings
x = 4  # Want figures to be A6
# plt.rc('figure', figsize=[46.82 * .5**(.5 * x), 33.11 * .5**(.5 * x)])
plt.rc('text', usetex=True)
plt.rc('font', family='serif')

# Define which colours you want to use
colors = ['blue', 'red', 'lightblue', 'pink']
# Define the groups
groups = ['High score of personality',
          'Low score of personality',]

dataFrames_dict = [personality_dataFrames[level]
                   for level in ['High', 'Low']]

# Get the max of the dataset
y_max = 6.5
# Get the min of the dataset
y_min = 0
# Calculate the y-axis range
y_range = y_max - y_min


# print("", len(x_pos_range))
# # Plot
df_template = None
for case in range(1):  # case 0, 1
    for personality in tqdm(personality_labels):  # personality 5
        # generate data for box plot
        datasets = []
        for idx, df_list in enumerate(dataFrames_dict):  # 4 each label in col
            # select cols
            start_index = 11 + case * 6
            df_qa = df_list[personality].iloc[:, start_index: start_index + 6]
            df_qa = df_qa.reset_index(drop=True)

            datasets.append(df_qa)

        assert len(datasets) == 2
        # # Create the plot
        fig = plt.figure(f"Case_{case}-{personality}",  figsize=(8, 6))

        ax = plt.axes()
        # Set x-positions for boxes
        x_pos_range = np.arange(len(datasets)) / (len(datasets) - 1)
        x_pos = (x_pos_range * 0.5) + 0.75

        for i, data in enumerate(datasets):
            positions = [x_pos[i] + j * 1 for j in range(len(data.T))]
            bp = ax.boxplot(
                np.array(data), sym='', whis=[0, 100], widths=0.6 / len(datasets),
                labels=list(x.split('with')[-1] for x in datasets[0]), patch_artist=True,
                positions=positions, showmeans=True,
            )
            # Fill the boxes with colours (requires patch_artist=True)
            k = i % len(colors)
            for box in bp['boxes']:
                box.set(facecolor=colors[k])
            # Make the median lines more visible
            plt.setp(bp['medians'], color='black')

            # Get the samples' medians
            medians = [bp['means'][j].get_ydata()[0]
                       for j in range(len(data.T))]
            medians = [str(round(s, 1)) for s in medians]
            # Increase the height of the plot by 5% to fit the labels
            ax.set_ylim([y_min - 0.1 * y_range, y_max + 0.05 * y_range])
            # Set the y-positions for the labels
            y_pos = y_min - 0.075 * y_range

            for tick, label in zip(range(len(data.T)), ax.get_xticklabels()):
                j = i % 2

                ax.text(
                    positions[tick], y_pos,
                    f'{medians[tick]}',
                    horizontalalignment='center', size='large',
                    color=colors[j],
                )
        # Axis details
        details = ax.set(
            title=f"Case {case} with participants with high / low scores of {personality}",
            ylabel='score'
        )

        ax.set_xlabel('Sentence sets')
        ax.tick_params(axis='x', bottom=False)

        xticks = ax.set_xticks(np.arange(len(list(datasets[0]))) + 1)
        xticks = ax.set_xticks(
            np.arange(len(list(datasets[0])) + 1) + 0.5, minor=True)
        xlim = ax.set_xlim([0.5, len(list(datasets[0])) + 0.5])
        # Legend
        legend_elements = []
        for i in range(len(datasets)):
            j = i % len(groups)
            k = i % len(colors)
            legend_elements.append(Patch(facecolor=colors[k], label=groups[j]))
        # ax.legend(handles=legend_elements, fontsize=8)
        plt.gca().legend(
            legend_elements, groups,
            fontsize=8, loc='upper left'
        )
        # Background
        # ax.set_facecolor('0.8')
        plt.grid(True, color='black', linestyle='-.', linewidth=0.2)
        # plt.show()
        # plt.tight_layout()

        os.makedirs('./figures/exp_box_images_avg/', exist_ok=True)
        plt.savefig(
            f'./figures/exp_box_images_avg/Case_{case}-{personality}.png', bbox_inches='tight')
        plt.close()

        # add to document
        r.add_text(f">>{personality} in Case_{case}:")
        r.add_picture(
            f"./figures/exp_box_images_avg/Case_{case}-{personality}.png")

document.save(
    './plot_document/box_plot_personality_avg_exp.docx')
